import PyPDF2
from docx import Document
import re
from docx import Document
import json
import json
from rdflib import Graph, Literal, Namespace, RDF, URIRef

## This code will convert pdf data to output.docx file format
def extract_text_from_pdf(pdf_path, start_page=0, end_page=2):
    text = ""
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Ensure end_page is not greater than the total number of pages
        if end_page is None or end_page > len(pdf_reader.pages):
            end_page = len(pdf_reader.pages)
        
        for page_number in range(start_page, end_page):
            page = pdf_reader.pages[page_number]
            text += page.extract_text()
    return text

# Function to write text to a Word document
def write_text_to_word(text, output_word_path):
    document = Document()
    document.add_paragraph(text)
    document.save(output_word_path)


# Example PDF file path (replace this with your actual PDF file path)
pdf_path = "Advanced-Database-Systems-and-Data-Warehouses.pdf"

# Extract text from the PDF
pdf_text = extract_text_from_pdf(pdf_path)
print(pdf_text)
# Example Word document path (replace this with your desired output path)
output_word_path = "output.docx"

# Write extracted text to Word document
write_text_to_word(pdf_text, output_word_path)

print(f"Text from PDF has been written to {output_word_path}")

## This code section will take output.docx file and convert it into JSON file
# Regular expressions to extract information
module_number_pattern = re.compile(r"  (\S+ [-\S]*)")
module_name_pattern = re.compile(r"Degree\s*(.+)")
contents_pattern = re.compile(r"Degree\s*(.+)")
credit_points_pattern = re.compile(r"Degree\s*(.+)")

# Function to extract information from a page
def extract_information(page_text):
    result = {}

    # Extract Module Number
    match_module_number = module_number_pattern.search(page_text)
    if match_module_number:
        result["Modulnummer"] = match_module_number.group(1)

    # Extract Module Name
    match_module_name = module_name_pattern.search(page_text)
    if match_module_name:
        result["Modulname"] = match_module_name.group(1)

    # Extract Contents
    match_contents = contents_pattern.search(page_text)
    if match_contents:
        contents = match_contents.group(1).strip()
        # Remove unwanted characters
        contents = contents.replace('\u2022', '').replace('\n', '')
        result["Inhalte und Qualifikationsziele"] = contents

    # Extract Credit Points
    match_credit_points = credit_points_pattern.search(page_text)
    if match_credit_points:
        result["Leistungspunkte und Noten"] = match_credit_points.group(1)

    return result

# Function to read text from a Word document
def read_text_from_word(word_path):
    doc = Document(word_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Example Word document path (replace this with your actual Word file path)
input_word_path = "output.docx"

# Read text from the Word document
word_text = read_text_from_word(input_word_path)

# Split the text into pages
pages = re.split(r"Modulnummer", word_text)[1:]

# Extract information from each page
result_list = [extract_information(page) for page in pages]

# Convert the list of dictionaries to JSON
json_data = json.dumps(result_list, indent=2)

# Write JSON data to a separate file
output_json_path = "output.json"
with open(output_json_path, "w") as json_file:
    json_file.write(json_data)

print(f"JSON data has been written to {output_json_path}")

# Load JSON data
data = json.loads(json_data)

# RDF Namespace
module_ns = URIRef("http://tuc.web.engineering/module#")

# Create RDF graph
g = Graph()

# Iterate over each module in the JSON data
for module in data:
    module_number = module.get("Modulnummer", "")
    module_name = module.get("Modulname", "")
    content = module.get("Inhalte und Qualifikationsziele", "")
    credit_points = module.get("Leistungspunkte und Noten", "")

    # RDF URI for the module
    module_uri = URIRef(f"http://tuc.web.engineering/module#{module_name.replace(' ', '_')}")

    # Add RDF triples for the module
    g.add((module_uri, RDF.type, module_ns))
    g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasModuleNumber"), Literal(module_number, datatype="http://www.w3.org/2001/XMLSchema#string")))
    g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasName"), Literal(module_name, datatype="http://www.w3.org/2001/XMLSchema#string")))
    g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasContent"), Literal(content, datatype="http://www.w3.org/2001/XMLSchema#string")))
    g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasCreditPoints"), Literal(credit_points, datatype="http://www.w3.org/2001/XMLSchema#integer")))

# Serialize RDF graph to RDF/XML format
rdf_output = g.serialize(format="xml")

# Save RDF data to a file with proper encoding
output_rdf_path = "output.rdf"
with open(output_rdf_path, "w", encoding="utf-8") as rdf_file:
    rdf_file.write(rdf_output)

print(f"RDF data has been saved to {output_rdf_path}")
